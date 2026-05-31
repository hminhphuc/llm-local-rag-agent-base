"""
Build TAI_LIEU_CHI_TIET.docx từ markdown source + Mermaid diagrams.

Quy trình:
1. Đọc TAI_LIEU_CHI_TIET.md
2. Inject diagram PNG vào đúng vị trí
3. Convert markdown -> docx qua pandoc
4. Post-process: thêm cover page, TOC, header/footer qua python-docx

Chạy:
    python docs/build_docx.py
Output:
    docs/TAI_LIEU_CHI_TIET.docx
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC_MD = REPO_ROOT / "TAI_LIEU_CHI_TIET.md"
DIAGRAMS_DIR = REPO_ROOT / "docs" / "diagrams"
SCREENSHOTS_DIR = REPO_ROOT / "docs" / "screenshots"
OUTPUT_DOCX = REPO_ROOT / "docs" / "TAI_LIEU_CHI_TIET.docx"
INTERMEDIATE_MD = REPO_ROOT / "docs" / "_intermediate.md"

# ============================================================
# Vị trí inject diagram - regex match phần heading, chèn ảnh ngay sau
# ============================================================
DIAGRAM_INJECTIONS = [
    # (Heading pattern để tìm vị trí, đường dẫn ảnh, caption)
    (
        r"^## Mục lục",
        "diagrams/01_overall_architecture.png",
        "Hình 1. Kiến trúc tổng thể: Local LLM + RAG + Agent (100% offline)",
        "before",  # chèn trước heading này
    ),
    (
        r"^## 2\.1 Embedding model lý thuyết",
        "diagrams/04_embedding_space.png",
        "Hình 2. Không gian embedding: text có nghĩa gần nhau → vector gần nhau",
        "after",
    ),
    (
        r"^## 2\.2 Chunking strategies chi tiết",
        "diagrams/02_rag_pipeline.png",
        "Hình 3. Pipeline RAG 6 bước: offline (build index) và online (query)",
        "before",
    ),
    (
        r"^## 3\.1 ReAct và các pattern khác",
        "diagrams/05_agent_equation.png",
        "Hình 4. Công thức Agent = LLM + Tools + ReAct Loop",
        "after",
    ),
    (
        r"^### ReAct \(Reason \+ Act\)",
        "diagrams/03_react_loop.png",
        "Hình 5. ReAct loop: Thought → Action → Observation, lặp đến khi đủ thông tin",
        "after",
    ),
    (
        r"^## 3\.4 MCP",
        "diagrams/06_mcp_architecture.png",
        "Hình 6. Kiến trúc MCP: tách tool thành server độc lập, mọi AI client đều dùng được",
        "after",
    ),
    (
        r"^## 4\.2 Architecture by scale",
        "diagrams/07_production_scales.png",
        "Hình 7. Tiến hóa kiến trúc theo scale: từ lab cá nhân đến toàn đơn vị",
        "after",
    ),
    (
        r"^## 2\.6 Bảo mật RAG in-depth",
        "diagrams/08_security_risks.png",
        "Hình 8. 4 rủi ro RAG + 4 rủi ro Agent — chiến lược defend khác nhau",
        "after",
    ),
    # Screenshots minh họa kết quả
    (
        r"^# Phần 1 — Local LLM với Ollama \(sâu\)",
        "screenshots/terminal_ollama_chat.png",
        "Hình 9. Ollama CLI: chat với Qwen3:4b qua command line, hoàn toàn offline",
        "after",
    ),
    (
        r"^## 2\.4 Prompt engineering cho RAG",
        "screenshots/terminal_rag_query.png",
        "Hình 10. Output thực tế của python rag_minimal.py: retrieve top-3 + sinh câu trả lời có trích nguồn",
        "after",
    ),
    (
        r"^## 3\.2 Tool calling mechanics",
        "screenshots/terminal_agent_trace.png",
        "Hình 11. Trace ReAct loop của agent: từng bước tool call → observation → final answer",
        "after",
    ),
    (
        r"^# Phần 2 — RAG đi sâu",
        "screenshots/open_webui_rag_real.png",
        "Hình 12. Open WebUI — giao diện chat khuyến nghị của workshop (giống ChatGPT, chạy 100% local): chọn model qwen3:1.7b, hỏi đáp tiếng Việt, hỗ trợ kéo–thả tài liệu để RAG",
        "after",
    ),
]


def inject_diagrams(md_text: str) -> str:
    """Chèn diagram image vào markdown source tại các vị trí đã định."""
    lines = md_text.split("\n")
    output = []
    used_imgs = set()

    for i, line in enumerate(lines):
        injected_before = False
        injected_after_payload = None

        for pattern, img, caption, position in DIAGRAM_INJECTIONS:
            if img in used_imgs:
                continue
            if re.match(pattern, line):
                img_md = f"\n![{caption}]({img})\n\n*{caption}*\n"
                if position == "before":
                    output.append(img_md)
                    injected_before = True
                else:
                    injected_after_payload = img_md
                used_imgs.add(img)
                break

        output.append(line)
        if injected_after_payload:
            output.append(injected_after_payload)

    return "\n".join(output)


def add_cover_and_postprocess(docx_path: Path) -> None:
    """Thêm cover page, header, footer, page numbers vào docx đã convert."""
    doc = Document(docx_path)

    # ===== Cover page (chèn ở đầu, trước nội dung hiện tại) =====
    # Cách python-docx: tạo paragraph mới ở đầu rồi style nó
    # Để chèn ở đầu, dùng XML manipulation

    body = doc.element.body

    # Tạo các paragraph cover (insert vào đầu)
    cover_elements = []

    def make_para(text: str, size: int, bold: bool, color_rgb: RGBColor,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after: int = 12):
        """Tạo 1 paragraph element XML."""
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        # alignment
        jc = OxmlElement("w:jc")
        align_map = {
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
        }
        jc.set(qn("w:val"), align_map.get(align, "left"))
        pPr.append(jc)

        # spacing
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), str(space_after * 20))
        pPr.append(spacing)
        p.append(pPr)

        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # font size (half-points)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(size * 2))
        rPr.append(sz)

        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)

        # color
        color = OxmlElement("w:color")
        color.set(qn("w:val"), f"{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}")
        rPr.append(color)

        # font name
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)

        r.append(rPr)

        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
        return p

    # Một số paragraph trống để đẩy text xuống giữa
    for _ in range(5):
        cover_elements.append(make_para("", 12, False, RGBColor(0, 0, 0)))

    cover_elements.append(make_para(
        "TÀI LIỆU CHI TIẾT",
        size=28, bold=True, color_rgb=RGBColor(0x1E, 0x40, 0xAF),
    ))
    cover_elements.append(make_para(
        "Workshop Local LLM + RAG + Agent",
        size=22, bold=False, color_rgb=RGBColor(0x1E, 0x40, 0xAF),
        space_after=24,
    ))
    cover_elements.append(make_para(
        "Chạy mô hình ngôn ngữ lớn local — Xây dựng RAG cho tài liệu nội bộ — Phát triển Agent biết dùng công cụ",
        size=13, bold=False, color_rgb=RGBColor(0x37, 0x37, 0x37),
        space_after=48,
    ))

    for _ in range(8):
        cover_elements.append(make_para("", 12, False, RGBColor(0, 0, 0)))

    cover_elements.append(make_para(
        "Dành cho học viên",
        size=14, bold=True, color_rgb=RGBColor(0x37, 0x37, 0x37),
    ))
    cover_elements.append(make_para(
        "Phiên bản 1.0 — 2026",
        size=12, bold=False, color_rgb=RGBColor(0x6B, 0x72, 0x80),
    ))

    # Page break sau cover
    pb_p = OxmlElement("w:p")
    pb_r = OxmlElement("w:r")
    pb_br = OxmlElement("w:br")
    pb_br.set(qn("w:type"), "page")
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    cover_elements.append(pb_p)

    # Insert vào đầu body (sau sectPr nếu có ở đầu, hoặc trước nội dung)
    first_real_child = None
    for child in body:
        if child.tag != qn("w:sectPr"):
            first_real_child = child
            break

    # addprevious inserts immediately before; iterate in order, each goes between previous and FRC
    for el in cover_elements:
        if first_real_child is not None:
            first_real_child.addprevious(el)
        else:
            body.append(el)

    # ===== Footer với page number =====
    for section in doc.sections:
        section.different_first_page_header_footer = True  # cover không có footer
        footer = section.footer
        # Clear footer existing content
        for p in footer.paragraphs:
            p.clear()

        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Trang X" với field code
        run = footer_para.add_run("Trang ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        r_element = run._element
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)

    # ===== Margins gọn hơn =====
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    doc.save(docx_path)
    print(f"[Post-process] Đã thêm cover, footer, margin: {docx_path}")


def main() -> None:
    print("[1/4] Đọc markdown nguồn...")
    md_text = SRC_MD.read_text(encoding="utf-8")
    print(f"      {len(md_text):,} ký tự, {md_text.count(chr(10)):,} dòng")

    print("[2/4] Inject diagram PNG vào markdown...")
    md_with_imgs = inject_diagrams(md_text)
    INTERMEDIATE_MD.write_text(md_with_imgs, encoding="utf-8")
    print(f"      Đã chèn {sum(1 for _ in DIAGRAM_INJECTIONS)} diagram")

    print("[3/4] Convert markdown -> docx qua pandoc...")
    # Reference doc giúp pandoc dùng style đẹp; nếu không có, dùng default
    reference_docx = REPO_ROOT / "docs" / "reference.docx"
    extra_args = [
        "--from=gfm+tex_math_dollars",
        "--to=docx",
        "--standalone",
        "--toc",
        "--toc-depth=3",
        f"--resource-path={REPO_ROOT / 'docs'}",
    ]
    if reference_docx.exists():
        extra_args.append(f"--reference-doc={reference_docx}")

    pypandoc.convert_file(
        str(INTERMEDIATE_MD),
        "docx",
        outputfile=str(OUTPUT_DOCX),
        extra_args=extra_args,
    )
    print(f"      Đã tạo {OUTPUT_DOCX}")

    print("[4/4] Post-process docx (cover page, footer, margins)...")
    add_cover_and_postprocess(OUTPUT_DOCX)

    # Cleanup intermediate
    INTERMEDIATE_MD.unlink(missing_ok=True)

    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"\n[OK] Hoàn tất: {OUTPUT_DOCX} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
